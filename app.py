import streamlit as st
import pytesseract
from PIL import Image
from docx import Document
from pdf2docx import Converter
import os
import io

# 1. Page Configuration (This gives Google your Title layout and description summary)
st.set_page_config(
    page_title="Universal Converter · Change is constant. Simple file tool.", 
    page_icon="📄", 
    layout="centered"
)

# 2. Hidden Google search verification data (Allows Google's indexing system to map the page)
st.markdown(
    '<meta name="google-site-verification" content="universal-converter-verification-tag" />', 
    unsafe_allow_html=True
)

# 3. Sidebar Navigation Panel
st.sidebar.header("Navigation Menu")
page = st.sidebar.radio("Go to:", ["Document Converter", "About This Tool", "Contact Support"])

# ==========================================
# TAB 1: THE ACTUAL WORKING CONVERTER
# ==========================================
if page == "Document Converter":
    st.title("📄 Universal Word Document Converter")
    st.write("Upload any PDF or image with text to instantly convert it into an editable Word file.")

    uploaded_file = st.file_uploader("Choose a file (PDF, PNG, JPG, JPEG, WEBP)", type=["pdf", "png", "jpg", "jpeg", "webp"])

    if uploaded_file is not None:
        file_name = uploaded_file.name
        base_name, extension = os.path.splitext(file_name)
        extension = extension.lower()
        
        st.info(f"🔄 Processing '{file_name}'... Please wait.")
        docx_buffer = io.BytesIO()
        success = False

        if extension == ".pdf":
            try:
                with open("temp_input.pdf", "wb") as f:
                    f.write(uploaded_file.getbuffer())
                cv = Converter("temp_input.pdf")
                cv.convert("temp_output.docx")
                cv.close()
                with open("temp_output.docx", "rb") as f:
                    docx_buffer.write(f.read())
                os.remove("temp_input.pdf")
                os.remove("temp_output.docx")
                success = True
            except Exception as e:
                st.error(f"❌ Failed to parse PDF: {e}")

        elif extension in [".png", ".jpg", ".jpeg", ".webp"]:
            try:
                img = Image.open(uploaded_file)
                extracted_text = pytesseract.image_to_string(img)
                doc = Document()
                has_content = False
                for line in extracted_text.split('\n'):
                    if line.strip():
                        doc.add_paragraph(line)
                        has_content = True
                if not has_content:
                    doc.add_paragraph("[AI Note: No text detected.]")
                doc.save(docx_buffer)
                success = True
            except Exception as e:
                st.error(f"❌ Failed image OCR: {e}")

        if success:
            st.success("🎉 Conversion Complete!")
            st.download_button(
                label="📥 Download Word Document",
                data=docx_buffer.getvalue(),
                file_name=f"{base_name}_converted.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

# ==========================================
# TAB 2: ABOUT PAGE
# ==========================================
elif page == "About This Tool":
    st.title("ℹ️ About This Project")
    st.write("This application was built to make file conversion simple and fast.")
    st.write("It automatically reads digital PDF formats and uses AI character recognition to read images.")

# ==========================================
# TAB 3: CONTACT PAGE
# ==========================================
elif page == "Contact Support":
    st.title("✉️ Get In Touch")
    st.write("Need help with a broken document? Leave a note below.")
    user_message = st.text_area("Enter your message or feedback:")
    if st.button("Submit Message"):
        if user_message:
            st.success("Thank you for your feedback!")
        else:
            st.warning("Please type a message first.")
ngml.document"
        )
